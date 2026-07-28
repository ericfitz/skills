"""Document object: holds the analyzed text, metadata, and cached spaCy doc.

The Document is the single object every check operates on. It carries:
    - The original raw bytes (for source_sha256)
    - The "extracted text" — what the analysis runs against. For markdown
      and plaintext this is essentially the source. For docx it's a
      reconstructed plaintext from python-docx.
    - The source_format flag, so locator-emitting code knows whether
      char/line offsets refer to the actual file or to extracted text.
    - A cached spaCy Doc (lazy-loaded on first access).
    - Paragraph and sentence indices, precomputed.
    - Heading structure (extracted at parse time), so section_path is
      available everywhere.

Extending:
    To support a new format, add a parser function and route it from
    `load_document`. The parser must return (extracted_text, headings)
    where headings is a list of (level, text, char_offset) tuples.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

# Module-level spaCy nlp object, loaded once.
_nlp: Any = None


def _get_nlp(model_name: str = "en_core_web_sm") -> Any:
    """Lazy-load spaCy. Cache at module level."""
    global _nlp
    if _nlp is None:
        import spacy  # imported here so checks that don't need it stay fast  # ty:ignore[unresolved-import]

        _nlp = spacy.load(model_name)
    return _nlp


@dataclass
class Heading:
    level: int
    text: str
    char_start: int
    char_end: int
    line_start: int


@dataclass
class Paragraph:
    index: int
    char_start: int
    char_end: int
    line_start: int
    line_end: int
    text: str


@dataclass
class SentenceRef:
    """Reference to a sentence in the document."""

    index: int
    char_start: int
    char_end: int
    paragraph_index: int
    text: str


@dataclass
class Document:
    """Loaded document ready for analysis.

    Most fields are populated by load_document(). Lazy fields (spacy_doc,
    sentences) are computed on first access.
    """

    source_path: str
    source_format: str  # "markdown" | "plaintext" | "docx" | "pdf"
    source_bytes: bytes
    source_sha256: str
    extracted_text: str
    extracted_text_sha256: str
    headings: list[Heading]
    paragraphs: list[Paragraph]

    # PDF-only side table: list of (char_start, char_end, page_number) tuples
    # mapping spans of extracted_text to their source page (1-indexed).
    # Empty for non-PDF formats.
    page_boundaries: list[tuple[int, int, int]] = field(default_factory=list)

    # Lazy-computed
    _spacy_doc: Any = None
    _sentences: list[SentenceRef] | None = None

    # Resolved via config (spaCy model + version pinning)
    _spacy_model_name: str = "en_core_web_sm"

    @property
    def spacy_doc(self) -> Any:
        """The spaCy-parsed document, computed once.

        spaCy's default `nlp.max_length` (1,000,000 characters) is too low
        for some long documents in our corpus (e.g., NIST SP 800-53 at
        ~1.65M chars). We raise it per-call to fit the document, with a
        safety ceiling of 10M to avoid OOM on malformed input.
        """
        if self._spacy_doc is None:
            nlp = _get_nlp(self._spacy_model_name)
            text_len = len(self.extracted_text)
            if text_len > nlp.max_length:
                if text_len > 10_000_000:
                    raise ValueError(
                        f"Document is {text_len:,} characters, exceeding the 10,000,000 "
                        "character safety ceiling. Refusing to parse."
                    )
                # Bump just enough to fit, with headroom.
                nlp.max_length = max(nlp.max_length, text_len + 100_000)
            self._spacy_doc = nlp(self.extracted_text)
        return self._spacy_doc

    @property
    def sentences(self) -> list[SentenceRef]:
        """All sentences with their offsets and paragraph membership."""
        if self._sentences is not None:
            return self._sentences
        sentences: list[SentenceRef] = []
        for idx, sent in enumerate(self.spacy_doc.sents):
            char_start = sent.start_char
            char_end = sent.end_char
            paragraph_index = self._paragraph_index_for_offset(char_start)
            sentences.append(
                SentenceRef(
                    index=idx,
                    char_start=char_start,
                    char_end=char_end,
                    paragraph_index=paragraph_index,
                    text=sent.text,
                )
            )
        self._sentences = sentences
        return sentences

    def _paragraph_index_for_offset(self, char_offset: int) -> int:
        """Binary-search-ish lookup; paragraphs are ordered by char_start."""
        for p in self.paragraphs:
            if p.char_start <= char_offset < p.char_end:
                return p.index
        # Fall back to last paragraph
        return self.paragraphs[-1].index if self.paragraphs else 0

    def page_number_for_offset(self, char_offset: int) -> int | None:
        """Return the 1-indexed page number for a char offset, or None for
        non-PDF documents (or if the offset falls outside any page range)."""
        for s, e, page in self.page_boundaries:
            if s <= char_offset < e:
                return page
        if self.page_boundaries and char_offset >= self.page_boundaries[-1][1]:
            return self.page_boundaries[-1][2]
        return None

    def section_path_for_offset(self, char_offset: int) -> list[str]:
        """Build the heading hierarchy in effect at a given char offset."""
        path: list[str] = []
        # Walk headings in order; for each, replace path at its level
        for h in self.headings:
            if h.char_start > char_offset:
                break
            # Trim path to this heading's level
            path = path[: h.level - 1]
            path.append(h.text)
        return path

    @property
    def char_count(self) -> int:
        return len(self.extracted_text)

    @property
    def line_count(self) -> int:
        return self.extracted_text.count("\n") + 1

    @property
    def word_count(self) -> int:
        # Use spaCy's tokenization; count alpha tokens
        return sum(1 for tok in self.spacy_doc if tok.is_alpha)


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------


def load_document(
    path: str | Path, spacy_model_name: str = "en_core_web_sm"
) -> Document:
    """Load a document from disk. Dispatches by extension."""
    path = Path(path)
    raw_bytes = path.read_bytes()
    source_sha = "sha256:" + hashlib.sha256(raw_bytes).hexdigest()
    suffix = path.suffix.lower()

    page_boundaries: list[tuple[int, int, int]] = []
    if suffix in (".md", ".markdown"):
        text, headings = _parse_markdown(raw_bytes)
        source_format = "markdown"
    elif suffix in (".txt", ".text", ""):
        text, headings = _parse_plaintext(raw_bytes)
        source_format = "plaintext"
    elif suffix == ".docx":
        text, headings = _parse_docx(path)
        source_format = "docx"
    elif suffix == ".pdf":
        text, headings, page_boundaries = _parse_pdf(path)
        source_format = "pdf"
    else:
        raise ValueError(f"Unsupported file extension: {suffix!r}")

    extracted_sha = "sha256:" + hashlib.sha256(text.encode("utf-8")).hexdigest()
    paragraphs = _segment_paragraphs(text)

    return Document(
        source_path=str(path),
        source_format=source_format,
        source_bytes=raw_bytes,
        source_sha256=source_sha,
        extracted_text=text,
        extracted_text_sha256=extracted_sha,
        headings=headings,
        paragraphs=paragraphs,
        page_boundaries=page_boundaries,
        _spacy_model_name=spacy_model_name,
    )


def _parse_markdown(raw_bytes: bytes) -> tuple[str, list[Heading]]:
    """Markdown: extract headings from ATX-style (`# `) lines.

    Treat the source markdown as the analyzed text — we keep heading
    syntax in place because removing it would shift offsets in confusing
    ways. The script is content-aware enough to skip heading prefixes
    when relevant.
    """
    text = raw_bytes.decode("utf-8")
    headings: list[Heading] = []
    line_offset = 0
    for line_idx, line in enumerate(text.split("\n"), start=1):
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            char_start = line_offset + len(m.group(1)) + 1  # after "# "
            char_end = line_offset + len(line)
            headings.append(
                Heading(
                    level=level,
                    text=heading_text,
                    char_start=char_start,
                    char_end=char_end,
                    line_start=line_idx,
                )
            )
        line_offset += len(line) + 1  # +1 for the newline
    return text, headings


def _parse_plaintext(raw_bytes: bytes) -> tuple[str, list[Heading]]:
    """Plain text: no heading extraction; treat as one flat document."""
    text = raw_bytes.decode("utf-8")
    return text, []


def _parse_docx(path: Path) -> tuple[str, list[Heading]]:
    """Extract text from .docx using python-docx.

    Builds a normalized plaintext: each paragraph on its own line,
    paragraphs separated by blank lines. Headings detected from
    paragraph styles (Heading 1, Heading 2, etc.).

    Note: char_start/line_start in returned Headings refer to the
    *extracted* text, not the .docx archive.
    """
    from docx import (  # ty:ignore[unresolved-import]
        Document as DocxDocument,  # python-docx  # ty:ignore[unresolved-import]
    )

    doc = DocxDocument(str(path))
    text_parts: list[str] = []
    headings: list[Heading] = []
    char_offset = 0
    line_offset = 1

    for para in doc.paragraphs:
        para_text = para.text or ""
        style_name = (para.style.name if para.style else "") or ""
        # Detect heading by style name pattern
        m = re.match(r"^Heading (\d+)$", style_name)
        if m and para_text.strip():
            level = int(m.group(1))
            headings.append(
                Heading(
                    level=level,
                    text=para_text.strip(),
                    char_start=char_offset,
                    char_end=char_offset + len(para_text),
                    line_start=line_offset,
                )
            )
        text_parts.append(para_text)
        char_offset += len(para_text) + 2  # paragraph + "\n\n"
        line_offset += para_text.count("\n") + 2

    text = "\n\n".join(text_parts)
    return text, headings


def _parse_pdf(path: Path) -> tuple[str, list[Heading], list[tuple[int, int, int]]]:
    """Extract text from a PDF using pypdf.

    Returns (extracted_text, headings, page_boundaries) where:
      - headings is always [] (we don't attempt heading detection from PDF
        layout; see docs/decisions.md)
      - page_boundaries is a list of (char_start, char_end, page_number)
        tuples mapping spans of the extracted_text back to their source
        page (page_number is 1-indexed)

    Pages are joined with two newlines so paragraph segmentation behaves
    sensibly across page breaks. Within a page we keep pypdf's extracted
    layout as-is.

    Raises ValueError if the PDF contains no extractable text (e.g., a
    scanned/image-only PDF). OCR is not supported.
    """
    from pypdf import PdfReader  # ty:ignore[unresolved-import]

    reader = PdfReader(str(path))
    page_texts: list[str] = []
    page_boundaries: list[tuple[int, int, int]] = []

    char_offset = 0
    page_separator = "\n\n"
    for page_idx, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        # Normalize line endings; pypdf occasionally emits \r line terminators.
        page_text = page_text.replace("\r\n", "\n").replace("\r", "\n")
        page_start = char_offset
        page_end = page_start + len(page_text)
        page_boundaries.append((page_start, page_end, page_idx))
        page_texts.append(page_text)
        # Advance offset by page text + separator (separator only between pages)
        char_offset = page_end
        if page_idx < len(reader.pages):
            char_offset += len(page_separator)

    text = page_separator.join(page_texts)

    if not any(c.isalpha() for c in text):
        raise ValueError(
            f"PDF contains no extractable text: {path!s}. "
            "Image-only / scanned PDFs require OCR, which is not supported."
        )

    return text, [], page_boundaries


def _segment_paragraphs(text: str) -> list[Paragraph]:
    """Segment by blank lines. Index from 0."""
    paragraphs: list[Paragraph] = []
    char_offset = 0
    line_offset = 1
    idx = 0
    # Split on one-or-more blank lines, keeping offsets
    chunks = re.split(r"(\n\s*\n)", text)
    for chunk in chunks:
        if not chunk:
            continue
        if re.match(r"^\n\s*\n$", chunk):
            char_offset += len(chunk)
            line_offset += chunk.count("\n")
            continue
        # This is a paragraph
        para_text = chunk
        p_char_start = char_offset
        p_char_end = char_offset + len(para_text)
        p_line_start = line_offset
        p_line_end = line_offset + para_text.count("\n")
        if para_text.strip():  # skip whitespace-only chunks
            paragraphs.append(
                Paragraph(
                    index=idx,
                    char_start=p_char_start,
                    char_end=p_char_end,
                    line_start=p_line_start,
                    line_end=p_line_end,
                    text=para_text,
                )
            )
            idx += 1
        char_offset += len(chunk)
        line_offset += chunk.count("\n")
    return paragraphs
